# Copyright (c) Microsoft Corporation. All rights reserved.
# Licensed under the MIT License.
#
# Generates the "Managed (Harness) Agents — Getting Started" Word document.
# PM-oriented framing for early-access customers.
#
# Run:
#   python generate_getting_started.py
#
# Output: managed-agents-getting-started.docx in this directory.

from __future__ import annotations

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def _ensure_code_style(doc: Document) -> None:
    if "Code Block" in [s.name for s in doc.styles]:
        return
    style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    pf = style.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.25)


def _shade(p, fill="F2F2F2") -> None:
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    _shade(p)
    p.add_run(text.rstrip("\n"))


def inline(p, text: str) -> None:
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    rem = text
    while rem:
        i = rem.find("[[c:")
        if i < 0:
            p.add_run(rem)
            break
        p.add_run(rem[:i])
        e = rem.find("]]", i)
        inline(p, rem[i + 4 : e])
        rem = rem[e + 2 :]


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        rem = it
        while rem:
            i = rem.find("[[c:")
            if i < 0:
                p.add_run(rem)
                break
            p.add_run(rem[:i])
            e = rem.find("]]", i)
            inline(p, rem[i + 4 : e])
            rem = rem[e + 2 :]


def h1(doc, t): doc.add_heading(t, level=1)
def h2(doc, t): doc.add_heading(t, level=2)


def table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for r, row in enumerate(rows, 1):
        for c, v in enumerate(row):
            cell = t.rows[r].cells[c]
            p = cell.paragraphs[0]
            rem = v
            while rem:
                i = rem.find("[[c:")
                if i < 0:
                    p.add_run(rem)
                    break
                p.add_run(rem[:i])
                e = rem.find("]]", i)
                inline(p, rem[i + 4 : e])
                rem = rem[e + 2 :]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def build() -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    _ensure_code_style(doc)

    title = doc.add_paragraph()
    r = title.add_run("Managed (Harness) Agents — Getting Started")
    r.bold = True
    r.font.size = Pt(24)
    sub = doc.add_paragraph()
    s = sub.add_run("Early-access guide · CLI and SDK · Microsoft Foundry")
    s.italic = True
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    meta = doc.add_paragraph()
    meta.add_run("Status: Preview    Audience: early-access customers    Updated: June 2026").italic = True
    doc.add_paragraph()

    h1(doc, "Why managed agents")
    para(doc,
         "A managed agent lets you ship a working AI agent by declaring just two things: a model and "
         "instructions. Microsoft Foundry provisions and runs the Brain+Hand sandbox for you — there is no "
         "container to build, no service code to host, and no infrastructure to manage. You go from idea to a "
         "deployed, callable agent in minutes.")
    bullets(doc, [
        "Time-to-first-agent measured in minutes, not days.",
        "No Dockerfile, no servers, no scaling decisions — the platform owns the runtime.",
        "One agent, two front doors: create with the CLI or the SDK; both target the same Foundry project.",
        "Standard OpenAI-shape Responses API for invocation, so existing tooling fits.",
    ])

    h1(doc, "What you'll need")
    bullets(doc, [
        "An Azure subscription and a Foundry project (a [[c:CognitiveServices/accounts/projects]] resource).",
        "A model deployment in that project (e.g. [[c:gpt-4.1-mini]]).",
        "Sign-in via [[c:azd auth login]] / [[c:az login]].",
        "Project endpoint [[c:AZURE_AI_PROJECT_ENDPOINT]] = https://<account>.services.ai.azure.com/api/projects/<project>",
        "Model name [[c:AZURE_AI_MODEL_DEPLOYMENT_NAME]] = e.g. gpt-4.1-mini",
    ])

    h1(doc, "Option A — azd CLI (fastest path)")
    h2(doc, "1. Install")
    code(doc,
         "winget install microsoft.azd\n"
         "azd extension install microsoft.azd.extensions\n"
         "azd extension source add --name MHA-dev --type url "
         "--location https://raw.githubusercontent.com/kshitij-microsoft/azure-dev/"
         "refs/heads/kchawla/azd-managed-harness/cli/azd/extensions/registry.json\n"
         "azd extension install azure.ai.agents --source MHA-dev\n"
         "azd auth login")
    h2(doc, "2. Create")
    para(doc, "Choose Prompt agent, pick your subscription and Foundry project, choose a model, and name it.")
    code(doc, "azd ai agent init")
    h2(doc, "3. Deploy and use")
    code(doc, "azd up\nazd ai agent list\nazd ai agent show\nazd ai agent invoke \"hello, what is your name?\"")
    para(doc, "`azd down` removes the agent with the project resources.")

    h1(doc, "Option B — Python SDK")
    h2(doc, "1. Install")
    code(doc,
         "pip install azure-ai-projects==2.3.0a20260625001 "
         "--extra-index-url https://pkgs.dev.azure.com/azure-sdk/public/_packaging/"
         "azure-sdk-for-python/pypi/simple\n"
         "pip install azure-identity python-dotenv")
    h2(doc, "2. Create a managed agent")
    code(doc,
         "from azure.identity import DefaultAzureCredential\n"
         "from azure.ai.projects import AIProjectClient\n"
         "from azure.ai.projects.models import PromptAgentDefinition, AgentHarness\n\n"
         "client = AIProjectClient(endpoint=endpoint, credential=DefaultAzureCredential(), allow_preview=True)\n\n"
         "client.agents.create_version(\n"
         "    agent_name=\"my-managed-agent\",\n"
         "    definition=PromptAgentDefinition(\n"
         "        model=model_name,\n"
         "        instructions=\"You are a helpful assistant.\",\n"
         "        harness=AgentHarness.GHCP,\n"
         "    ),\n"
         ")")
    h2(doc, "3. Invoke")
    code(doc,
         "openai_client = client.get_openai_client()\n"
         "response = openai_client.responses.create(\n"
         "    input=[{\"role\": \"user\", \"content\": \"Generate python to print the OS and run it.\"}],\n"
         "    store=False,\n"
         "    extra_body={\"agent_reference\": {\"name\": \"my-managed-agent\", \"version\": \"1\", "
         "\"type\": \"agent_reference\"}},\n"
         ")")

    h1(doc, "What a response looks like")
    para(doc, "Invocations stream Server-Sent Events from the project data-plane. The Brain plans the turn and "
              "the Hand sandbox runs any tools/code; only [[c:output_text.delta]] events carry visible text.")
    code(doc,
         "POST .../api/projects/<project>/openai/v1/responses\n"
         "x-agent-session-id: ses_...\n\n"
         "event: response.created\n"
         "event: response.output_text.delta\n"
         "event: response.completed")

    h1(doc, "CLI vs SDK at a glance")
    table(doc,
          ["Task", "CLI", "SDK"],
          [
              ["Create", "azd ai agent init + azd up", "create_version(..., harness=GHCP)"],
              ["Invoke", "azd ai agent invoke", "responses.create(agent_reference)"],
              ["List / show", "azd ai agent list / show", "agents.list / get_version"],
              ["Tear down", "azd down", "delete on the project"],
          ])

    h1(doc, "Recommended next steps")
    bullets(doc, [
        "Pick the path that matches the customer: CLI for hands-on demos, SDK for app integration.",
        "Both write to the same Foundry project — an agent made via SDK shows up in the CLI and vice versa.",
        "Share feedback on time-to-first-agent and any blocking errors during the bug bash.",
    ])
    return doc


def main() -> None:
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "managed-agents-getting-started.docx")
    build().save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
