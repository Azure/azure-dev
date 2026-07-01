# Copyright (c) Microsoft Corporation. All rights reserved.
# Licensed under the MIT License.
#
# Generates the "Managed (Harness) Agents in azd" Word document spec.
#
# Run:
#   python generate_spec.py
#
# Output: spec.docx in the same directory.

from __future__ import annotations

import os
from dataclasses import dataclass

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# ----------------------------- styling helpers -----------------------------


def _ensure_code_style(doc: Document) -> None:
    """Create a "Code Block" character style (Consolas, dark gray)."""
    styles = doc.styles
    if "Code Block" in [s.name for s in styles]:
        return
    style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Consolas"
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    pf = style.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.25)


def _ensure_inline_code_style(doc: Document) -> None:
    styles = doc.styles
    if "InlineCode" in [s.name for s in styles]:
        return
    style = styles.add_style("InlineCode", WD_STYLE_TYPE.CHARACTER)
    style.font.name = "Consolas"
    style.font.size = Pt(10)


def _shade_paragraph(paragraph, fill_hex: str = "F2F2F2") -> None:
    """Apply a background fill to a paragraph (for code blocks)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def add_code_block(doc: Document, code: str, language: str | None = None) -> None:
    para = doc.add_paragraph(style="Code Block")
    _shade_paragraph(para)
    para.add_run(code.rstrip("\n"))


def add_inline_code(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_para(doc: Document, text: str) -> None:
    """Add a normal paragraph. Use [[code:foo]] to render inline code spans."""
    para = doc.add_paragraph()
    remaining = text
    while remaining:
        idx = remaining.find("[[code:")
        if idx == -1:
            para.add_run(remaining)
            break
        para.add_run(remaining[:idx])
        end = remaining.find("]]", idx)
        if end == -1:
            para.add_run(remaining[idx:])
            break
        add_inline_code(para, remaining[idx + 7 : end])
        remaining = remaining[end + 2 :]


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        remaining = item
        while remaining:
            idx = remaining.find("[[code:")
            if idx == -1:
                para.add_run(remaining)
                break
            para.add_run(remaining[:idx])
            end = remaining.find("]]", idx)
            if end == -1:
                para.add_run(remaining[idx:])
                break
            add_inline_code(para, remaining[idx + 7 : end])
            remaining = remaining[end + 2 :]


def add_h1(doc: Document, text: str) -> None:
    para = doc.add_heading(text, level=1)
    para.paragraph_format.space_before = Pt(18)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            para = cell.paragraphs[0]
            # Honor inline code in cells.
            remaining = val
            while remaining:
                idx = remaining.find("[[code:")
                if idx == -1:
                    para.add_run(remaining)
                    break
                para.add_run(remaining[:idx])
                end = remaining.find("]]", idx)
                if end == -1:
                    para.add_run(remaining[idx:])
                    break
                add_inline_code(para, remaining[idx + 7 : end])
                remaining = remaining[end + 2 :]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


# ----------------------------- content builders ---------------------------


def build_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Managed (Harness) Agents in azd")
    run.bold = True
    run.font.size = Pt(26)

    subtitle = doc.add_paragraph()
    sub = subtitle.add_run(
        "Design spec for the `managed` agent kind in the `azd ai agent` extension"
    )
    sub.italic = True
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    meta = doc.add_paragraph()
    meta.add_run("Status: Implemented (preview)    Owner: azd ai agent team    Audience: azd contributors").italic = True

    doc.add_paragraph()  # spacer


def build_overview(doc: Document) -> None:
    add_h1(doc, "Overview")
    add_para(
        doc,
        "The Azure AI Foundry agent platform exposes three first-class agent kinds today: "
        "[[code:hosted]] (bring-your-own container/code), [[code:workflow]] (multi-step orchestration), "
        "and a new [[code:managed]] kind backed by the Prompt Execution Service (PES) Brain+Hand harness. "
        "Managed agents are the simplest shape: the customer declares a model deployment and "
        "system instructions; the platform provisions the runtime, executes turns, and persists "
        "state. There is no container to build, no Dockerfile, and no service code on the customer side.",
    )
    add_para(
        doc,
        "This spec describes how the [[code:azure.ai.agents]] azd extension implements first-class "
        "support for managed agents end-to-end: YAML schema, API wire types, ARM-shaped HTTP client, "
        "init scaffolding flow, delete dispatch, and local-development affordances against the "
        "[[code:managed-harness]] vienna backend.",
    )


def build_goals(doc: Document) -> None:
    add_h1(doc, "Goals and Non-Goals")
    add_h3(doc, "Goals")
    add_bullets(
        doc,
        [
            "Add a [[code:managed]] discriminator to [[code:AgentKind]] and an accompanying "
            "[[code:ManagedAgent]] YAML type that can round-trip through the existing parser.",
            "Map the YAML definition to the wire shape ([[code:ManagedAgentDefinition]] + "
            "[[code:ManagedEnvironment]] + [[code:ManagedPackages]]) accepted by the v2.0 "
            "managed-agents controller.",
            "Add an ARM-rooted HTTP client ([[code:ManagedAgentClient]]) covering the lifecycle "
            "(create / get / update / delete / list) and the Responses subtree "
            "(create / get / cancel / delete).",
            "Wire the init flow to ask which agent kind to create as the very first interactive "
            "step, and add a [[code:runInitManaged]] path that scaffolds the minimum surface "
            "(agent.yaml + azure.yaml service entry) with no Docker, no Language, no src/.",
            "Wire the delete flow to detect managed agents via the YAML discriminator and route "
            "deletion through [[code:ManagedAgentClient.DeleteAgent]] rather than the hosted path.",
            "Allow developer machines to target a local [[code:managed-harness]] backend "
            "without an Azure login (env-var override + credential-skip for localhost).",
            "Keep all existing hosted-agent code paths byte-identical when [[code:kind]] is not "
            "[[code:managed]].",
        ],
    )

    add_h3(doc, "Non-Goals (this milestone)")
    add_bullets(
        doc,
        [
            "[[code:azd ai agent show]] / [[code:list]] / [[code:invoke]] wiring for managed agents "
            "(designed but deferred — see Open Questions).",
            "Hosted versioning semantics for managed agents — the backend does not expose a per-version "
            "delete on the v2.0 surface, so [[code:--version]] is rejected with a typed validation error.",
            "Surfacing every advanced [[code:ManagedAgentDefinition]] field "
            "([[code:structured_inputs]], [[code:files]], full [[code:environment]] block) "
            "through YAML — only [[code:model]], [[code:instructions]], [[code:skills]], "
            "and [[code:policies]] are exposed today.",
            "Automatic ARM workspace discovery from a Foundry project endpoint — callers must "
            "set [[code:AZD_MANAGED_AGENT_SUBSCRIPTION_ID]] / [[code:_RESOURCE_GROUP]] / "
            "[[code:_WORKSPACE]] explicitly.",
            "Schema publication — the [[code:agent.yaml]] schema annotation points at "
            "[[code:microsoft/AgentSchema]] and assumes that repo will pick up a "
            "[[code:ManagedAgent.yaml]] sibling alongside the existing kinds.",
        ],
    )


def build_user_stories(doc: Document) -> None:
    add_h1(doc, "User Stories")
    add_bullets(
        doc,
        [
            "As a developer I run [[code:azd ai agent init]] in an empty folder, choose "
            "\u201cManaged agent\u201d, answer three prompts (name / model / instructions), "
            "and end up with an [[code:agent.yaml]] and an [[code:azure.yaml]] service entry "
            "ready for [[code:azd deploy]].",
            "As a developer I set [[code:FOUNDRY_PROJECT_ENDPOINT]] in my azd environment, run "
            "[[code:azd deploy]], and the managed agent is created on Foundry without azd "
            "building any container or pushing any code.",
            "As a developer I run [[code:azd ai agent delete --service <name>]] and the "
            "extension detects [[code:kind: managed]] in the service's [[code:agent.yaml]] and "
            "deletes via the managed lifecycle endpoint instead of the hosted one.",
            "As a Foundry platform contributor I run a local [[code:managed-harness]] vienna "
            "backend on [[code:http://localhost:5000]], set [[code:AZD_FOUNDRY_ENDPOINT_OVERRIDE=1]] "
            "and [[code:AZD_MANAGED_AGENT_BASE_URL=http://localhost:5000]], and exercise the full "
            "azd-managed-agent surface against my dev box without an Azure login.",
        ],
    )


def build_architecture(doc: Document) -> None:
    add_h1(doc, "Architecture")
    add_para(
        doc,
        "Managed support is layered onto the existing extension along the same seams as the other "
        "agent kinds. The discriminator is [[code:agent.yaml \u2192 kind]]; everything downstream "
        "switches on that value.",
    )
    add_code_block(
        doc,
        """\
azure.ai.agents extension
\u251c\u2500 internal/pkg/agents/agent_yaml/
\u2502    \u251c\u2500 yaml.go              \u2190 ManagedAgent struct + AgentKindManaged constant
\u2502    \u251c\u2500 parse.go             \u2190 switch on kind \u2192 unmarshal to ManagedAgent
\u2502    \u251c\u2500 map.go               \u2190 CreateManagedAgentAPIRequest(...) \u2192 wire type
\u2502    \u2514\u2500 managed_test.go      \u2190 round-trip / validate / dispatcher coverage
\u2502
\u251c\u2500 internal/pkg/agents/agent_api/
\u2502    \u251c\u2500 models.go            \u2190 ManagedAgentDefinition / ManagedEnvironment / ManagedPackages
\u2502    \u251c\u2500 managed_operations.go\u2190 ManagedAgentClient (lifecycle + responses) + BuildWorkspaceRoutePrefix
\u2502    \u2514\u2500 managed_operations_test.go
\u2502
\u2514\u2500 internal/cmd/
     \u251c\u2500 init.go                 \u2190 prompts kind first; routes to runInitManaged when "managed"
     \u251c\u2500 init_from_templates_helpers.go \u2190 promptAgentKind() Select
     \u251c\u2500 init_managed.go         \u2190 scaffolds agent.yaml + adds azure.yaml service entry (no Docker, no src/)
     \u251c\u2500 managed_dispatch.go     \u2190 isManagedAgentYAML / newManagedAgentClientFromEnv / localhost detection
     \u251c\u2500 delete.go               \u2190 detects kind \u2192 runManagedDelete via ManagedAgentClient
     \u2514\u2500 project_endpoint.go     \u2190 AZD_FOUNDRY_ENDPOINT_OVERRIDE bypass for local http:// targets
""",
    )
    add_para(
        doc,
        "The split between [[code:agent_yaml]] and [[code:agent_api]] mirrors the hosted/workflow "
        "kinds: [[code:agent_yaml]] is the customer-authored shape, [[code:agent_api]] is the "
        "wire shape sent to Foundry. [[code:agent_yaml/map.go]] is the only crossover.",
    )


def build_yaml_schema(doc: Document) -> None:
    add_h1(doc, "YAML Schema")
    add_para(
        doc,
        "A managed [[code:agent.yaml]] is small by design \u2014 the platform owns the runtime, "
        "so the customer-authored shape is just [[code:kind]], [[code:name]], [[code:model]], "
        "[[code:instructions]], optional [[code:skills]], and optional [[code:policies]].",
    )

    add_h3(doc, "Minimum example")
    add_code_block(
        doc,
        """\
# yaml-language-server: $schema=https://raw.githubusercontent.com/microsoft/AgentSchema/refs/heads/main/schemas/v1.0/ManagedAgent.yaml

kind: managed
name: customer-support-bot
model: gpt-4.1-mini
instructions: |
  You are a helpful customer-support agent.
  Always reply in the user's language and cite a knowledge-base
  article when you give a factual answer.
""",
    )

    add_h3(doc, "Full example (skills + RAI policy)")
    add_code_block(
        doc,
        """\
# yaml-language-server: $schema=https://raw.githubusercontent.com/microsoft/AgentSchema/refs/heads/main/schemas/v1.0/ManagedAgent.yaml

kind: managed
name: research-assistant
displayName: Research Assistant
description: Summarizes long-form web content with citations.
model: gpt-4.1-mini
instructions: |
  You are a research assistant. Cite every source you use.
skills:
  - foundry.tools.web_search
  - foundry.tools.code_interpreter
policies:
  - type: rai_policy
    rai_policy_name: /subscriptions/<sub>/resourceGroups/<rg>/providers/Microsoft.CognitiveServices/accounts/<acct>/raiPolicies/<policy>
""",
    )

    add_h3(doc, "Field reference")
    add_table(
        doc,
        ["Field", "Required", "Type", "Notes"],
        [
            ["[[code:kind]]", "yes", "string", "Must be the literal [[code:managed]]."],
            ["[[code:name]]", "yes", "string", "Foundry agent identity. Also used as the folder name when scaffolding into a non-empty cwd."],
            ["[[code:displayName]]", "no", "string", "Optional human-readable label."],
            ["[[code:description]]", "no", "string", "Optional description."],
            ["[[code:metadata]]", "no", "map", "Free-form key/value tags. [[code:authors]] is special-cased into a comma-separated string by the mapper."],
            ["[[code:model]]", "yes", "string", "Model deployment name (e.g. [[code:gpt-4.1-mini]]). Validated non-empty by [[code:CreateManagedAgentAPIRequest]]."],
            ["[[code:instructions]]", "yes", "string", "System/developer message inserted into the model context. Validated non-empty by [[code:CreateManagedAgentAPIRequest]]."],
            ["[[code:skills]]", "no", "string[]", "Optional list of Foundry skill identifiers attached to the agent."],
            ["[[code:policies]]", "no", "Policy[]", "Optional governance policies. Today only [[code:type: rai_policy]] with an ARM-id-shaped [[code:rai_policy_name]] is supported."],
        ],
    )

    add_h3(doc, "Discriminator routing")
    add_para(
        doc,
        "[[code:agent_yaml/parse.go]] switches on the [[code:kind]] field and unmarshals into "
        "the matching type. The managed branch is symmetric with [[code:hosted]] and [[code:workflow]]:",
    )
    add_code_block(
        doc,
        """\
switch agentDef.Kind {
case AgentKindHosted:
    // ... ContainerAgent
case AgentKindWorkflow:
    // ... Workflow
case AgentKindManaged:
    var agent ManagedAgent
    if err := yaml.Unmarshal(data, &agent); err != nil {
        return nil, fmt.Errorf("failed to unmarshal to ManagedAgent: %w", err)
    }
    return agent, nil
}
return nil, fmt.Errorf("unrecognized agent kind: %s", agentDef.Kind)
""",
    )


def build_wire_contract(doc: Document) -> None:
    add_h1(doc, "API Wire Contract")
    add_para(
        doc,
        "The wire shape lives in [[code:internal/pkg/agents/agent_api/models.go]]. It is the "
        "JSON body POSTed under the standard [[code:CreateAgentRequest]] envelope, with "
        "[[code:Definition]] set to a [[code:ManagedAgentDefinition]].",
    )

    add_h3(doc, "[[code:ManagedAgentDefinition]]")
    add_code_block(
        doc,
        """\
type ManagedAgentDefinition struct {
    AgentDefinition
    Model            string              `json:"model"`
    Instructions     string              `json:"instructions,omitempty"`
    Tools            []any               `json:"tools,omitempty"`
    ToolChoice       any                 `json:"tool_choice,omitempty"`
    Skills           []string            `json:"skills,omitempty"`
    StructuredInputs map[string]any      `json:"structured_inputs,omitempty"`
    Environment      *ManagedEnvironment `json:"environment,omitempty"`
    Files            map[string]string   `json:"files,omitempty"`
}
""",
    )

    add_h3(doc, "[[code:ManagedEnvironment]] / [[code:ManagedPackages]]")
    add_code_block(
        doc,
        """\
type ManagedPackages struct {
    Pip []string `json:"pip,omitempty"`
    Apt []string `json:"apt,omitempty"`
}

type ManagedEnvironment struct {
    BaseImage            *string           `json:"base_image,omitempty"`
    Image                *string           `json:"image,omitempty"`
    Packages             *ManagedPackages  `json:"packages,omitempty"`
    CPU                  *string           `json:"cpu,omitempty"`
    Memory               *string           `json:"memory,omitempty"`
    EgressPolicy         *string           `json:"egress_policy,omitempty"`
    EnvironmentVariables map[string]string `json:"environment_variables,omitempty"`
}
""",
    )

    add_h3(doc, "Mapping rules ([[code:CreateManagedAgentAPIRequest]])")
    add_bullets(
        doc,
        [
            "[[code:model]] and [[code:instructions]] are validated non-empty; otherwise the call "
            "returns [[code:fmt.Errorf]] (\u201cmanaged agent requires a non-empty model/instructions\u201d).",
            "[[code:policies]] are run through [[code:mapRaiConfig]] (the same helper used by hosted "
            "agents) to produce [[code:AgentDefinition.RaiConfig]] on the wire.",
            "[[code:skills]] are cloned ([[code:append([]string(nil), ...]]) so the request does not "
            "alias the customer-supplied slice.",
            "When a non-nil [[code:AgentBuildConfig]] carries [[code:EnvironmentVariables]], they are "
            "copied (via [[code:maps.Clone]]) into [[code:Environment.EnvironmentVariables]] so the "
            "Hand sandbox can read them.",
            "No [[code:image]], [[code:cpu]], [[code:memory]], or [[code:endpoint]] fields are set "
            "from the YAML \u2014 these belong to the hosted/container shape and are not part of the "
            "managed customer-authored surface today.",
        ],
    )


def build_url_surface(doc: Document) -> None:
    add_h1(doc, "URL Surface and [[code:ManagedAgentClient]]")
    add_para(
        doc,
        "All managed operations are rooted at an ARM-shaped workspace resource. The client takes a "
        "[[code:BaseURL]] (origin) and a [[code:RoutePrefix]] (everything between the origin and "
        "[[code:/agents]]) so the same client targets production, an alternate cloud, or a local "
        "[[code:managed-harness]] backend without rewiring URL assembly:",
    )
    add_code_block(
        doc,
        """\
{baseURL}{routePrefix}/agents
{baseURL}{routePrefix}/agents/{name}
{baseURL}{routePrefix}/agents/{name}/openai/responses
{baseURL}{routePrefix}/agents/{name}/openai/responses/{responseId}
{baseURL}{routePrefix}/agents/{name}/openai/responses/{responseId}/cancel
""",
    )

    add_h3(doc, "Production route prefix")
    add_para(
        doc,
        "Built by [[code:BuildWorkspaceRoutePrefix(sub, rg, ws)]]. Each segment is "
        "[[code:url.PathEscape]]'d:",
    )
    add_code_block(
        doc,
        """\
/agents/v2.0/subscriptions/<sub>/resourceGroups/<rg>/providers/Microsoft.MachineLearningServices/workspaces/<ws>
""",
    )

    add_h3(doc, "Operations")
    add_table(
        doc,
        ["Method", "URL suffix (after [[code:{baseURL}{routePrefix}/agents]])", "Accepted statuses", "Notes"],
        [
            ["[[code:CreateAgent]]", "", "200, 201", "POST. Body is [[code:CreateAgentRequest]] (envelope) with a [[code:ManagedAgentDefinition]]."],
            ["[[code:GetAgent]]", "/{name}", "200", "Path-escaped agent name."],
            ["[[code:UpdateAgent]]", "/{name}", "200", "POST (intentional; the controller treats POST-to-name as replace)."],
            ["[[code:DeleteAgent]]", "/{name}?force=<bool>", "200, 204", "Accepts 204 because vienna returns it in some configs; synthesizes [[code:{Deleted: true, Name: name}]] when the body is empty."],
            ["[[code:ListAgents]]", "?kind/limit/after/before/order", "200", "All filters optional."],
            ["[[code:CreateResponse]]", "/{name}/openai/responses", "200, 201, 202", "Body passes through verbatim \u2014 callers serialize the OpenAI Responses shape themselves; caller-supplied headers are forwarded."],
            ["[[code:GetResponse]]", "/{name}/openai/responses/{responseId}", "200", "Raw body + cloned response headers returned."],
            ["[[code:CancelResponse]]", "/{name}/openai/responses/{responseId}/cancel", "200, 202", "POST."],
            ["[[code:DeleteResponse]]", "/{name}/openai/responses/{responseId}", "200, 204", ""],
        ],
    )

    add_h3(doc, "Construction")
    add_code_block(
        doc,
        """\
client, err := agent_api.NewManagedAgentClient(agent_api.ManagedAgentClientOptions{
    BaseURL:     "https://management.azure.com",
    RoutePrefix: prefix,    // from BuildWorkspaceRoutePrefix(sub, rg, ws)
    Credential:  cred,      // azcore.TokenCredential (nil for unauthenticated localhost)
    Scopes:      nil,       // defaults to {"https://ai.azure.com/.default"}
})
""",
    )

    add_h3(doc, "Pipeline policies")
    add_bullets(
        doc,
        [
            "[[code:NewMsCorrelationPolicy]] \u2014 attaches [[code:x-ms-correlation-request-id]].",
            "[[code:NewUserAgentPolicy]] \u2014 [[code:azd-ext-azure-ai-agents/<version>]].",
            "[[code:NewBearerTokenPolicy]] (when [[code:Credential]] is non-nil) using scopes "
            "[[code:https://ai.azure.com/.default]] by default. The policy is prepended so it runs "
            "before correlation/user-agent.",
            "Logging is configured with [[code:IncludeBody=true]] and an allowlist for "
            "[[code:X-Ms-Correlation-Request-Id]] / [[code:X-Request-Id]].",
        ],
    )

    add_h3(doc, "API version")
    add_para(
        doc,
        "Sent on every request via the [[code:api-version]] query param. The default is "
        "[[code:DefaultManagedAgentAPIVersion = \"2025-08-01-preview\"]]. Defining it as an "
        "exported constant keeps test wire assertions and command call sites in lockstep when "
        "the backend rolls forward.",
    )


def build_cli_surface(doc: Document) -> None:
    add_h1(doc, "CLI Surface")

    add_h3(doc, "[[code:azd ai agent init]] \u2014 kind selection")
    add_para(
        doc,
        "Before any hosted-specific detection runs (manifest discovery, [[code:--src]] handling, "
        "deploy-mode/runtime/entry-point flags), the [[code:init]] command asks the user which "
        "kind to scaffold. The prompt is suppressed when any \u201chosted signal\u201d is present "
        "on the command line so existing CI scripts stay on the hosted path:",
    )
    add_code_block(
        doc,
        """\
hostedSignalsPresent := userProvidedManifest ||
    flags.src != "" ||
    flags.deployMode != "" ||
    flags.runtime != "" ||
    flags.entryPoint != ""
if !hostedSignalsPresent {
    kindChoice, kindErr := promptAgentKind(ctx, azdClient, flags.noPrompt)
    if kindErr != nil {
        return kindErr
    }
    if kindChoice == AgentKindChoiceManaged {
        return runInitManaged(ctx, flags, azdClient)
    }
}
""",
    )
    add_para(
        doc,
        "[[code:promptAgentKind]] returns [[code:AgentKindChoiceHosted]] in [[code:--no-prompt]] "
        "mode to preserve today's behaviour for callers that do not yet know about the new kind.",
    )

    add_h3(doc, "[[code:runInitManaged]] \u2014 scaffolding flow")
    add_bullets(
        doc,
        [
            "Prompt for [[code:agent name]] (default [[code:my-managed-agent]]; "
            "[[code:--agent-name]] required in [[code:--no-prompt]] mode).",
            "Prompt for [[code:model deployment]] (default [[code:gpt-4.1-mini]]; "
            "[[code:--model]] required in [[code:--no-prompt]] mode).",
            "Prompt for [[code:system instructions]] (default placeholder; in [[code:--no-prompt]] "
            "mode a self-documenting stub is written so the customer can edit before deploying).",
            "Resolve target directory: write at the cwd when empty, otherwise create a sanitized "
            "subfolder named after the agent. Refuses to clobber a non-empty existing subfolder.",
            "Write [[code:agent.yaml]] with the [[code:yaml-language-server]] schema annotation "
            "pointing at the [[code:ManagedAgent.yaml]] schema.",
            "Add an [[code:azure.yaml]] service entry via [[code:azdClient.Project().AddService]] "
            "with [[code:Host: azure.ai.agent]] and no [[code:Language]] / no [[code:Docker]]. "
            "If no [[code:azure.yaml]] exists, return a typed dependency error pointing the user "
            "at [[code:azd init]] \u2014 we intentionally do not scaffold a project here.",
            "Print a concise summary and a copy-paste-able next-steps block "
            "([[code:azd env set FOUNDRY_PROJECT_ENDPOINT \u2026]] \u2192 [[code:azd deploy]]).",
        ],
    )
    add_para(
        doc,
        "Critically, [[code:runInitManaged]] does NOT call [[code:ensureProject]] / clone a "
        "[[code:azd-ai-starter-basic]] template / scaffold any Bicep. Managed agents assume the "
        "Foundry project endpoint already exists \u2014 forcing the user through hosted-shaped "
        "infra scaffolding would be misleading.",
    )

    add_h3(doc, "[[code:azd ai agent delete]] \u2014 dispatch")
    add_para(
        doc,
        "The delete command inspects the service's [[code:agent.yaml]] via "
        "[[code:isManagedAgentYAML]]. If the discriminator is [[code:managed]], it routes to "
        "[[code:DeleteAction.runManagedDelete]]:",
    )
    add_bullets(
        doc,
        [
            "[[code:--version]] is rejected up-front with a typed validation error "
            "([[code:CodeInvalidParameter]]) \u2014 managed agents do not expose per-version delete "
            "on the v2.0 surface.",
            "Constructs a [[code:ManagedAgentClient]] via [[code:newManagedAgentClientFromEnv]].",
            "Calls [[code:DeleteAgent(ctx, name, DefaultManagedAgentAPIVersion, force)]] and "
            "passes [[code:--force]] through to the [[code:force]] query parameter.",
            "On success, best-effort cleans up the matching env-var keys and session state to "
            "stay at parity with the hosted delete path.",
            "Honors [[code:--output json]] by emitting [[code:DeleteAgentResponse]] directly; "
            "the default human-readable output is a one-liner ([[code:Managed agent \"<name>\" deleted.]]).",
        ],
    )


def build_dispatch_and_envvars(doc: Document) -> None:
    add_h1(doc, "Lifecycle Dispatch and Environment Variables")
    add_para(
        doc,
        "All command-side managed plumbing lives in [[code:internal/cmd/managed_dispatch.go]]. "
        "It provides three helpers plus a small block of env-var constants:",
    )

    add_h3(doc, "[[code:isManagedAgentYAML(filePath string) (bool, error)]]")
    add_para(
        doc,
        "Reads the file and unmarshals only the [[code:kind]] field via a probe struct. A missing "
        "file returns [[code:(false, nil)]] so callers can treat \u201cno agent.yaml present\u201d "
        "as \u201cnot a managed agent\u201d rather than an error. A malformed file returns a wrapped "
        "[[code:yaml.Unmarshal]] error so the surrounding command can surface it.",
    )

    add_h3(doc, "[[code:newManagedAgentClientFromEnv(ctx) (*ManagedAgentClient, error)]]")
    add_bullets(
        doc,
        [
            "Reads [[code:AZD_MANAGED_AGENT_SUBSCRIPTION_ID]] / [[code:AZD_MANAGED_AGENT_RESOURCE_GROUP]] / "
            "[[code:AZD_MANAGED_AGENT_WORKSPACE]] from the process environment.",
            "When any of the three is missing or empty, returns a typed validation error "
            "([[code:CodeInvalidParameter]]) whose suggestion lists exactly which env vars to set.",
            "Builds the route prefix via [[code:BuildWorkspaceRoutePrefix]] (so the same escaping "
            "and validation rules apply as the unit tests in [[code:managed_operations_test.go]]).",
            "Resolves the base URL from [[code:AZD_MANAGED_AGENT_BASE_URL]] when set, "
            "otherwise falls back to [[code:https://management.azure.com]].",
            "Resolves the credential via [[code:newAgentCredentialOrNil]]: nil for localhost targets "
            "(so devs do not need an Azure login to talk to a local backend), otherwise the standard "
            "agent credential. Credential-construction failures are intentionally swallowed and "
            "surfaced as nil so the underlying HTTP 401/403 becomes the user-visible error \u2014 "
            "that error is more actionable than a generic \u201cfailed to create credential\u201d wrap.",
        ],
    )

    add_h3(doc, "[[code:isLocalBackendBaseURL(baseURL)]]")
    add_para(
        doc,
        "Hostname-only prefix check for [[code:localhost]], [[code:127.0.0.1]], and [[code:[::1]]] "
        "with either [[code:http://]] or [[code:https://]]. Non-standard ports still match. Used "
        "by both the credential decision above and the [[code:project_endpoint]] validator bypass.",
    )

    add_h3(doc, "Environment variable reference")
    add_table(
        doc,
        ["Variable", "Type", "Required", "Used by", "Notes"],
        [
            ["[[code:AZD_MANAGED_AGENT_SUBSCRIPTION_ID]]", "string", "yes (managed ops)", "[[code:newManagedAgentClientFromEnv]]", "Azure subscription id hosting the workspace."],
            ["[[code:AZD_MANAGED_AGENT_RESOURCE_GROUP]]", "string", "yes (managed ops)", "[[code:newManagedAgentClientFromEnv]]", "Resource group containing the workspace."],
            ["[[code:AZD_MANAGED_AGENT_WORKSPACE]]", "string", "yes (managed ops)", "[[code:newManagedAgentClientFromEnv]]", "Workspace name. Path-escaped before being placed in the route prefix."],
            ["[[code:AZD_MANAGED_AGENT_BASE_URL]]", "string", "no", "[[code:newManagedAgentClientFromEnv]]", "Overrides the default [[code:https://management.azure.com]] origin. Used to point at a local [[code:managed-harness]] dev backend ([[code:http://localhost:5000]])."],
            ["[[code:AZD_FOUNDRY_ENDPOINT_OVERRIDE]]", "presence", "no", "[[code:project_endpoint]] validator", "When set to any non-empty value, the validator accepts [[code:http://]] in addition to [[code:https://]] and skips the Foundry host-suffix check. Intentionally undocumented in user help \u2014 dev/test only."],
            ["[[code:FOUNDRY_PROJECT_ENDPOINT]]", "string", "yes (deploy)", "azd environment", "Foundry project endpoint used at deploy/invoke time. Unchanged from the hosted flow \u2014 listed here only because the [[code:runInitManaged]] next-steps block prints how to set it."],
        ],
    )


def build_local_dev(doc: Document) -> None:
    add_h1(doc, "Local Development Against the [[code:managed-harness]] Backend")
    add_para(
        doc,
        "The vienna [[code:managed-harness]] service implements the same v2.0 controller as "
        "production and is the recommended local backend. To target it from azd:",
    )
    add_code_block(
        doc,
        """\
# 1) start managed-harness locally
#    (default port 5000; see vienna repo for details)

# 2) bypass the strict project-endpoint validator so http://localhost is accepted
$Env:AZD_FOUNDRY_ENDPOINT_OVERRIDE = "1"

# 3) point the managed client at the local origin (anything in {localhost,127.0.0.1,::1} works)
$Env:AZD_MANAGED_AGENT_BASE_URL = "http://localhost:5000"

# 4) supply the ARM workspace tuple (the harness validates shape but does not call ARM)
$Env:AZD_MANAGED_AGENT_SUBSCRIPTION_ID = "00000000-0000-0000-0000-000000000000"
$Env:AZD_MANAGED_AGENT_RESOURCE_GROUP  = "local-rg"
$Env:AZD_MANAGED_AGENT_WORKSPACE       = "local-ws"

# 5) scaffold a managed agent and deploy
azd ai agent init      # choose "Managed agent"
azd env set FOUNDRY_PROJECT_ENDPOINT http://localhost:5000/api/projects/local
azd deploy
""",
    )
    add_para(
        doc,
        "Because [[code:isLocalBackendBaseURL]] returns true for any of these origins, the "
        "[[code:ManagedAgentClient]] is constructed without a credential and the bearer-token "
        "policy is omitted from the pipeline \u2014 no Azure login is required.",
    )


def build_testing(doc: Document) -> None:
    add_h1(doc, "Testing Strategy")
    add_bullets(
        doc,
        [
            "[[code:agent_yaml/managed_test.go]] \u2014 round-trip YAML \u2192 [[code:ManagedAgent]] "
            "\u2192 YAML; validator coverage for missing [[code:model]] / [[code:instructions]]; "
            "discriminator routing through [[code:parse.go]].",
            "[[code:agent_yaml/map_test.go]] \u2014 [[code:CreateManagedAgentAPIRequest]] populates "
            "[[code:ManagedAgentDefinition]] correctly, copies skills/policies, and propagates "
            "build-time env vars into [[code:ManagedEnvironment]].",
            "[[code:agent_api/managed_operations_test.go]] \u2014 [[code:BuildWorkspaceRoutePrefix]] "
            "input validation; [[code:httptest]]-backed coverage for every lifecycle and responses "
            "URL (including the 204 path on [[code:DeleteAgent]] and the [[code:force]] query param); "
            "header forwarding on [[code:CreateResponse]]; credential-nil pipeline construction.",
            "[[code:cmd/project_endpoint_test.go]] \u2014 the [[code:AZD_FOUNDRY_ENDPOINT_OVERRIDE]] "
            "bypass accepts [[code:http://]] and skips the host-suffix check only when set.",
            "[[code:cmd/managed_dispatch_test.go]] \u2014 [[code:isManagedAgentYAML]] handles "
            "missing/malformed/non-managed/managed files; [[code:newManagedAgentClientFromEnv]] "
            "returns typed validation errors with actionable suggestions when env vars are missing.",
        ],
    )


def build_open_questions(doc: Document) -> None:
    add_h1(doc, "Open Questions and Future Work")
    add_bullets(
        doc,
        [
            "Wire [[code:azd ai agent show]] / [[code:list]] / [[code:invoke]] for managed agents. "
            "Designed but deferred this milestone. [[code:invoke]] in particular wants to lean on "
            "[[code:CreateResponse]] + [[code:GetResponse]] streaming.",
            "Surface advanced [[code:ManagedAgentDefinition]] fields through YAML: "
            "[[code:structured_inputs]], [[code:files]], and the full [[code:environment]] block "
            "(image, base_image, packages, cpu/memory, egress_policy). The wire types already accept "
            "them; only the [[code:ManagedAgent]] YAML struct intentionally omits them today.",
            "Automatic ARM workspace discovery from a Foundry project endpoint. Today the user "
            "must set three env vars explicitly. A future iteration could derive the workspace tuple "
            "from a project endpoint + credential via a lightweight Foundry control-plane lookup.",
            "Hosted versioning parity. Whether the v2.0 controller will gain a per-version delete is "
            "an open product question. Today [[code:--version]] is a typed validation error.",
            "Publish [[code:ManagedAgent.yaml]] in [[code:microsoft/AgentSchema]] so the schema URL "
            "in the [[code:yaml-language-server]] annotation resolves to a real document.",
            "Telemetry: emit a [[code:azd.ext.azure.ai.agents.kind]] field on init/delete events "
            "so we can measure managed adoption distinct from hosted.",
        ],
    )


def build_references(doc: Document) -> None:
    add_h1(doc, "References")
    add_bullets(
        doc,
        [
            "[[code:cli/azd/extensions/azure.ai.agents/internal/pkg/agents/agent_yaml/yaml.go]] "
            "\u2014 [[code:ManagedAgent]] struct, [[code:AgentKindManaged]] constant.",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/pkg/agents/agent_yaml/parse.go]] "
            "\u2014 discriminator switch.",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/pkg/agents/agent_yaml/map.go]] "
            "\u2014 [[code:CreateManagedAgentAPIRequest]].",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/pkg/agents/agent_api/models.go]] "
            "\u2014 [[code:ManagedAgentDefinition]] / [[code:ManagedEnvironment]] / "
            "[[code:ManagedPackages]].",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/pkg/agents/agent_api/managed_operations.go]] "
            "\u2014 [[code:ManagedAgentClient]] + [[code:BuildWorkspaceRoutePrefix]].",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/cmd/init.go]] \u2014 kind prompt insertion site.",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/cmd/init_from_templates_helpers.go]] "
            "\u2014 [[code:promptAgentKind]].",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/cmd/init_managed.go]] "
            "\u2014 [[code:runInitManaged]] scaffolding.",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/cmd/managed_dispatch.go]] "
            "\u2014 dispatch helpers + env-var constants.",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/cmd/delete.go]] "
            "\u2014 [[code:runManagedDelete]].",
            "[[code:cli/azd/extensions/azure.ai.agents/internal/cmd/project_endpoint.go]] "
            "\u2014 [[code:AZD_FOUNDRY_ENDPOINT_OVERRIDE]] bypass.",
        ],
    )


# ----------------------------- assemble -----------------------------


def build_doc() -> Document:
    doc = Document()

    # Defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _ensure_code_style(doc)
    _ensure_inline_code_style(doc)

    build_title(doc)
    build_overview(doc)
    build_goals(doc)
    build_user_stories(doc)
    build_architecture(doc)
    build_yaml_schema(doc)
    build_wire_contract(doc)
    build_url_surface(doc)
    build_cli_surface(doc)
    build_dispatch_and_envvars(doc)
    build_local_dev(doc)
    build_testing(doc)
    build_open_questions(doc)
    build_references(doc)

    return doc


def main() -> None:
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "spec.docx")
    doc = build_doc()
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
